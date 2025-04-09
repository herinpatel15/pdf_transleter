from pdf2docx import Converter
from docx import Document
from googletrans import Translator
import os
import time
import comtypes.client

def convert_pdf_to_docx(pdf_file, docx_file):
    cv = Converter(pdf_file)
    cv.convert(docx_file, start=0, end=None)
    cv.close()
    print("✅ Converted PDF to DOCX.")

def translate_docx(input_docx, output_docx, target_lang='gu'):
    translator = Translator()
    doc = Document(input_docx)
    new_doc = Document()

    total_paragraphs = len(doc.paragraphs)
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            try:
                translated = translator.translate(text, dest=target_lang).text
                new_doc.add_paragraph(translated)
                time.sleep(1)  # helps avoid rate-limiting
            except Exception as e:
                print("❌ Translation failed:", e)
                new_doc.add_paragraph("[Translation Error]")
        else:
            new_doc.add_paragraph("")  # preserve blank lines

        # Show progress percentage
        percent = int(((idx + 1) / total_paragraphs) * 100)
        print(f"🔄 Translating: {percent}% complete", end='\r')  # overwrite the same line

    new_doc.save(output_docx)
    print("\n✅ Translated and saved DOCX.")

# def translate_docx(input_docx, output_docx, target_lang='gu'):
#     translator = Translator()
#     doc = Document(input_docx)
#     new_doc = Document()

#     for para in doc.paragraphs:
#         text = para.text.strip()
#         if text:
#             try:
#                 translated = translator.translate(text, dest=target_lang).text
#                 new_doc.add_paragraph(translated)
#                 time.sleep(1)
#             except Exception as e:
#                 print("❌ Translation failed:", e)
#                 new_doc.add_paragraph("[Translation Error]")
#         else:
#             new_doc.add_paragraph("")  # preserve blank lines

#     new_doc.save(output_docx)
#     print("✅ Translated and saved DOCX.")

def convert_docx_to_pdf(input_docx, output_pdf):
    # This part works only on Windows with MS Word installed
    
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(os.path.abspath(input_docx))
    doc.SaveAs(os.path.abspath(output_pdf), FileFormat=17)  # 17 = PDF
    doc.Close()
    word.Quit()
    print("✅ Converted DOCX to PDF.")

def pdf_to_translated_pdf(pdf_path, output_pdf_path, temp_docx='temp.docx', translated_docx='translated.docx', lang='gu'):
    convert_pdf_to_docx(pdf_path, temp_docx)
    print("translating...")
    translate_docx(temp_docx, translated_docx, target_lang=lang)
    print("converting docx to pdf...")
    convert_docx_to_pdf(translated_docx, output_pdf_path)

    # Cleanup if needed
    os.remove(temp_docx)
    os.remove(translated_docx)

# Example usage:
pdf_to_translated_pdf("./input_pdf_file/Unit 3_Developmental Psychology_10520403.pdf", "Unit 3_Developmental Psychology_10520403.pdf", lang="gu")

# from pdf2docx import Converter
# from docx import Document
# from googletrans import Translator
# import os
# import time
# import comtypes.client
# from docx.shared import Pt
# from docx.enum.style import WD_STYLE_TYPE

# def convert_pdf_to_docx(pdf_file, docx_file):
#     cv = Converter(pdf_file)
#     cv.convert(docx_file, start=0, end=None)
#     cv.close()
#     print("✅ Converted PDF to DOCX.")

# def translate_docx(input_docx, output_docx, target_lang='gu'):
#     translator = Translator()
#     doc = Document(input_docx)
#     new_doc = Document()

#     for para in doc.paragraphs:
#         text = para.text.strip()
#         style_name = para.style.name  # preserve original paragraph style

#         if text:
#             try:
#                 translated_text = translator.translate(text, dest=target_lang).text
#                 new_para = new_doc.add_paragraph()
#                 new_para.style = style_name if style_name in [s.name for s in new_doc.styles] else 'Normal'

#                 for run in para.runs:
#                     run_text = run.text.strip()
#                     if run_text:
#                         try:
#                             translated_run = translator.translate(run_text, dest=target_lang).text
#                         except Exception as e:
#                             print("Run translation error:", e)
#                             translated_run = "[Translation Failed]"

#                         new_run = new_para.add_run(translated_run)
#                         # Copy formatting
#                         new_run.bold = run.bold
#                         new_run.italic = run.italic
#                         new_run.underline = run.underline
#                         new_run.font.size = run.font.size
#                         new_run.font.name = run.font.name
#                         # Optionally more formatting...
#                         time.sleep(1)
#             except Exception as e:
#                 print("❌ Paragraph translation failed:", e)
#                 new_doc.add_paragraph("[Translation Error]")
#         else:
#             new_doc.add_paragraph("")

#     new_doc.save(output_docx)
#     print("✅ Translated and saved DOCX with formatting.")

# def convert_docx_to_pdf(input_docx, output_pdf):
#     word = comtypes.client.CreateObject('Word.Application')
#     doc = word.Documents.Open(os.path.abspath(input_docx))
#     doc.SaveAs(os.path.abspath(output_pdf), FileFormat=17)  # 17 = PDF
#     doc.Close()
#     word.Quit()
#     print("✅ Converted DOCX to PDF.")

# def pdf_to_translated_pdf(pdf_path, output_pdf_path, temp_docx='temp.docx', translated_docx='translated.docx', lang='gu'):
#     convert_pdf_to_docx(pdf_path, temp_docx)
#     print("translating...")
#     translate_docx(temp_docx, translated_docx, target_lang=lang)
#     print("converting docx to pdf...")
#     convert_docx_to_pdf(translated_docx, output_pdf_path)

#     # Cleanup temp files
#     os.remove(temp_docx)
#     os.remove(translated_docx)

# # Example usage:
# pdf_to_translated_pdf("sample.pdf", "translated_final_gujarati.pdf", lang="gu")
