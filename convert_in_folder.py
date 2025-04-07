from pdf2docx import Converter
from docx import Document
from googletrans import Translator
import os
import time
import comtypes.client

def convert_pdf_to_docx(pdf_file, docx_file):
    cv = Converter(pdf_file)
    cv.convert(docx_file, start=0, end=None)
    cv.close()
    print(f"✅ Converted {pdf_file} to DOCX.")

def translate_docx(input_docx, output_docx, target_lang='gu'):
    translator = Translator()
    doc = Document(input_docx)
    new_doc = Document()

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            try:
                translated = translator.translate(text, dest=target_lang).text
                new_doc.add_paragraph(translated)
                time.sleep(1)
            except Exception as e:
                print("❌ Translation failed:", e)
                new_doc.add_paragraph("[Translation Error]")
        else:
            new_doc.add_paragraph("")  # preserve blank lines

    new_doc.save(output_docx)
    print(f"✅ Translated and saved DOCX: {output_docx}")

def convert_docx_to_pdf(input_docx, output_pdf):
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(os.path.abspath(input_docx))
    doc.SaveAs(os.path.abspath(output_pdf), FileFormat=17)  # 17 = PDF
    doc.Close()
    word.Quit()
    print(f"✅ Converted DOCX to PDF: {output_pdf}")

def pdf_to_translated_pdf(pdf_path, output_pdf_path, temp_docx='temp.docx', translated_docx='translated.docx', lang='gu'):
    convert_pdf_to_docx(pdf_path, temp_docx)
    print("Translating...")
    translate_docx(temp_docx, translated_docx, target_lang=lang)
    print("Converting DOCX to PDF...")
    convert_docx_to_pdf(translated_docx, output_pdf_path)

    # Cleanup
    os.remove(temp_docx)
    os.remove(translated_docx)

def process_folder(folder_path, out_path, lang='gu'):
    for filename in os.listdir(folder_path):
        if filename.lower().endswith(".pdf"):
            pdf_path = os.path.join(folder_path, filename)
            name_only = os.path.splitext(filename)[0]
            output_pdf = os.path.join(out_path, f"translated_{name_only}.pdf")

            print(f"Processing {filename}...")
            pdf_to_translated_pdf(pdf_path, output_pdf, lang=lang)

            # Ask before processing each file
            # choice = input(f"\n➡️ Convert this file? {filename} (y/n): ").strip().lower()
            # if choice == 'y':
            #     try:
            #         pdf_to_translated_pdf(pdf_path, output_pdf, lang=lang)
            #     except Exception as e:
            #         print(f"❌ Failed to process {filename}: {e}")
            # else:
            #     print(f"⏩ Skipping {filename}.")

# 🟢 Example usage:
# Change the folder path below to your folder containing PDFs
folder_with_pdfs = "C:/hello_herin/pdf_transleter/input_pdf_file"  # change this to your actual folder
folder_output_pdfs = "C:/hello_herin/pdf_transleter/output_pdf_file"  # change this to your actual folder
process_folder(folder_with_pdfs, folder_output_pdfs, lang="gu")
